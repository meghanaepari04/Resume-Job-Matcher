# backend/utils.py
import re
import os
import io
from PyPDF2 import PdfReader
from docx import Document

# Basic skill list — expand this to cover your domain and target JD keywords.
SKILLS = [
    # Programming Languages
    "python", "java", "c", "c++", "c#", "javascript", "typescript", "php", "ruby", "go", "swift", "kotlin", "r", "scala",

    # Web & Frameworks
    "html", "css", "react", "angular", "vue", "next.js", "node", "express", "django", "flask", "spring boot", "fastapi",

    # Databases
    "mysql", "postgresql", "mongodb", "oracle", "sqlite", "redis", "dynamodb",

    # Cloud & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "terraform", "ansible", "ci/cd", "linux",

    # Data Science & AI
    "machine learning", "deep learning", "nlp", "pandas", "numpy", "scipy",
    "matplotlib", "seaborn", "plotly", "tensorflow", "keras", "pytorch", "scikit-learn",

    # Other Tools
    "git", "github", "gitlab", "jira", "confluence", "api", "graphql", "rest",

    # Soft Skills (optional, for HR-oriented matching)
    "communication", "leadership", "teamwork", "problem solving", "time management"
]


# Normalize skills to lower-case and strip spaces for matching
SKILLS = [s.lower() for s in SKILLS]

def extract_text_from_pdf(path_or_bytes):
    """
    Accepts a file path or bytes-like object and returns extracted text.
    """
    text = ""
    if isinstance(path_or_bytes, (bytes, bytearray)):
        reader = PdfReader(io.BytesIO(path_or_bytes))
    else:
        reader = PdfReader(path_or_bytes)
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(path_or_bytes):
    if isinstance(path_or_bytes, (bytes, bytearray)):
        # python-docx doesn't accept bytes directly; write to temp file if needed
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(path_or_bytes)
            tmp.flush()
            doc = Document(tmp.name)
    else:
        doc = Document(path_or_bytes)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def extract_text(file_storage):
    """
    file_storage is a Flask FileStorage object from request.files['resume'].
    Returns extracted text for PDF or DOCX.
    """
    filename = file_storage.filename.lower()
    data = file_storage.read()
    # Reset file pointer in case caller wants to reuse file_storage
    try:
        file_storage.stream.seek(0)
    except Exception:
        pass

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(data)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(data)
    else:
        # Try both methods fallback
        try:
            return extract_text_from_pdf(data)
        except Exception:
            return extract_text_from_docx(data)

def clean_text(text):
    """
    Minimal cleaning: lower, remove multiple spaces and non-ASCII artifacts.
    """
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    text = text.strip()
    return text

def extract_skills(text):
    """
    Simple keyword matching for skills. Returns a set of found skills.
    """
    text = clean_text(text)
    found = set()
    for skill in SKILLS:
        # use word boundary for accurate matches
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text):
            found.add(skill)
    return found

def compare_skills(resume_text, jd_text):
    resume_skills = [s for s in SKILLS if s.lower() in resume_text.lower()]
    jd_skills = [s for s in SKILLS if s.lower() in jd_text.lower()]
    matched =list(set(resume_skills) & set(jd_skills))
    missing = list(set(jd_skills) - set(resume_skills))
    extra = list(set(resume_skills) - set(jd_skills))
    return {
        "resume_skills": resume_skills,
        "jd_skills": jd_skills,
        "matched": matched,
        "missing": missing,
        "extra": extra
    }
